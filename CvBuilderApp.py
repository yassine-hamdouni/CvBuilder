import streamlit as st
import io
import json
import base64
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Logo (base64 embedded) ──────────────────────────────────────────────────
_LOGO_B64 = "/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCABEAOUDASIAAhEBAxEB/8QAHAABAAIDAQEBAAAAAAAAAAAAAAcIBAUGAwEC/8QAQhAAAQMDAwEGAwUDCAsAAAAAAQIDBAAFEQYHEiEIEzFBUYEUIpEVMmFxoRYjUhc0QmKSorHRRlNVcnN1gpOywtL/xAAaAQEAAwEBAQAAAAAAAAAAAAAAAQMEAgUG/8QALxEAAQQBAwIEBQMFAAAAAAAAAQACAxEEEiExBVETFCJBYXGBofAGscEjJEJSkf/aAAwDAQACEQMRAD8AuXSlKIlKUoiUpSiJSlKIlKUoiUpSiJSlKIlKUoiUpSiJSlKIlKUoiUpSiJSlKIlKUoiUpSiJSlY851bMZx1DSnloQpSW0+KyBkAfifCiL4/LYjgGQ80yD4FxYTn615/alu/2hE/76f8AOqY6v0hu9q29v3i9aUvr7zyypKFIBQynPRCBnCQB06VrIm0W5EqQhhGjbg2VnHN7ghCfxKiele0zpMOm3zAH6LzDnvuhGVeaJJYkpUqO864kHBU2sKGfasioAQ/dtgNt7c2iwsX1qS+pdylNyi0GX144pxxOUYASFeo/EV2Gze7Nt15AuKpEdFqnQf3j0db/ACSGf9aFEDIHUHp09685+I8MMjBbAata2ZDS4Mds7spPpVe4HaFmXnVyLDprRouBkyixDcXOKC4nPRwjgeIxlR9BW93f3nkbeaoZsg061cS5DRILvxZbwVFQ4gcT0+Xx/GuvI5HiCPT6iLr4KPNw6S69gpnpVeNQ9pOBCiQkWmw/HTXY6HZQXJ4ssLUAS2DjKyM4JwBXRbT762fWV2RY7lb1We6PdI4L3eMvkdeKVYBCvQEdfWpd0/KbGZHMNIMyFztIduplpXH7i69sOhbOLje3XMuZRGjtDk7IV6JHoPMnoKhJ3tPyvjf3WkGfh89ErnnvMeycVGPg5GQ3VG2x+d1MuVFEacd1Z2lcDtXujp/cGI79nJdi3COkGRCfI5pB8FJI6KTnpke4FcBuFv6/pbW9z04nSzcpMJ4NB8zijnkA5xwOPH18q4ZiTvkMYb6gpdkxNYHk7FT7XP3nWOl7PPXBu2obXBlJAUWZEpKFBJ8CQT4GotuW/wBalayh6bsFtVc23JCWHZ63uDeSepbABKgPU4zVdd1tWK13q9+/uW9EBTjCGO6DveY4AjPLA8c+lbMPpUs76kGkVazT57GN9G5V+IrzUmM1IYcS406gLQtJyFJIyCD6Yr0qFOzvui7rGSNMKsyISLVa2yHxJKy5xKW/u8Rj18TXztUa5uWltOQbTZJq4c+6uL5vNHDjbKAOXE+RJIGfzrOMGXzIx+CVo8yzwvF9lNlKrF2StQ3IXa6N3i+SHIUpTceKzKfUvvJRClkIKicHgkk+vSup7WOsbhYNNW6z2mc9DlXV5anXGVlKwwgDIBHUZUoDp5A1LsF4yhjA7n3/ADsuW5bTD4pCnSlVp7LmsrhC0/qebqi8OrsNrQ04l6U4pZaWrllCSck5AHy+pGPGvS+dp1lE1SbHpVb0YKwh2ZK7tSh68Eg4+tdHpuQZXRMF17rkZsWgPcatWSpUNbX782DVl2Zs9zhO2W4vniwFuhxl5R8EheBhXoCOvrWLu1vZN0FrFyyHSyJzXw7chl8zC2VpUDnpwPgoEeNVeRyPE8LT6uys81Fo13spvpXEas11Gse1Z1wI4kIVEZfZY73j3inOPFHLB/i8ceVcDtbvz+2GuIWnZWnkW0S0uBt5MsufOlJUE44jxwaiPEmkjdI1uzeVLsiNrg0ncqdaVEO9u8H8nV6g2xqyJuTkqMqQsqk913Y5cQOiTnOD9KkfRdzlXrSdru8yEmE/NjIkKjhfPu+YyBnAz0I8qrdBI2NsjhseF22VjnlgO4W3pSlVKxaHWl+Y03pqZe5CC4iKjIQDjmokBI9yRVbLhulrWZMXIF9dipUflajpSlCR6dQSfzNWtkR2JDRakMtutnxQtIUD7GsFdosyQVKtUAADJJjo6fpXn5mLLOQGP0j8+K9rpXUsXCDjNCHk9/b6UVVf+UfW5/0pnfVP+VSNsnc9daivhn3O7znrNHQoK70AIecPQJBx1x4nH4VxOtt+WWb5Ii6T0rYnYLDhQmVMj5U9g45JSnHFPpnJrXt9pHWjTYbRZtPIbSOiQw4AP79bsL9K9QJEj3kjtZ++6nqX6vwZoXRRY4aTtdDb5bBWn1BZ4N9s8m0XJhD8OW0Wn21D7ySP0PmD5GqH6tt69I6xvNktl3VIbjOOQzIZWR3rR+8hWPHp8qh4ZBqe7rvNqCHsuL9eGIkG/Xl5xmztsJUnDIABkEKJOB1x5E8ai7aHau4bhWa93YSlx0x2yiE4vqJEr7xCif6OOhPqoehr6bpcZw2PfOabdV3PdfF5rhkOa2Mb8qVeyNomHHsDutX3WX58wrjxwkhXwzSThQPotRHUeQx61w3bC6bpRf8AlTX/AJuVq9idfSNudaPWW/lyNbJL/cT2XOhivJPEOY/DwV6jr5Ctn2v1oXudCcQpK0KtLSkqScgjm51FX4zHjqhc42CNj8Fw8t8npAog7qTtltptEyNtbZPutliXSbc4yZD78gFRHPqEo6/KAMDp1zk1Xjce0t6I3PucC0vr4Wual2KsqypAAS4kE+oyBn8K6ewbm7jbb6ciWVLcX7PkR0yLc7LYKwltwch3agQCOp+U5wc1p9ttK37c7XgkSRIkMOyhIutwWn5QnOVDPhyVjASP8BU47JYJJJp32w37/wAfZcSuZK1kcbfV8lu+1TdJdw3NQy6rDUW2sd0jySXE81H3J/Suhseo9IRNEs6ee2Uv01tccJelGIlTryyOrgXx5A56jB6dK3far25uVykx9Y2GE5LDUYR58dlPJaUJzwcSkdSACQcegNcHC391pC0anTbYhCQ2x8M3PIUH20AYHTOOYHTl74zVcP8AcYkQiF1zvVLuT+lO8v2v4WvDYm2ams+7thkpst3jRnJCmHluxHEp7pSVZCiRj09xWq7QAKt5dSpAGTKSBn/hoqV+zbatxLzfG9Sakvl/FkYQruGZkpwiW4RgHio9UJBJyfE4qKN/VpG9eoskdJiM9f6iKvx5tWe6yLDa2+YVckdYw5on+FYLQuwGkrKm13Kc/Pm3WMpD5c7/AINd4OuAgDqnPqTmoJ7S9ptli3SmQLPAjwIqYTKwywgJSFFJJOB5mrtR8FhsjzSP8Kqt2u9JXX9r29Vxobz9tkRUMPutoKgy4gkAKx4AgjB8OhrzOlZb35dyu5B5WzOx2tg9DVO+1ml9O2bTtqutss0KHNlWxjv32WglbmUJUcnz69arH2o799tbtTI7a+TFqaRDQP6w+Zf95WPapF7Nm52o9QTZGnbmiJIi2m0d4wppspdX3ZSlIJzg9OnhVemUTtV6rQ2eS514n4PmebrnX6ZP0rT0/GMWVI+U3pHe+VTlzNfAxrByf2W81EZ+nNLaPZYL8N9xDl8S+Bj94tfFsg+fFDSf7X41sN9dUz9UauiPXJCGpEO2R2XW0H5UPKQHHMf9Sse1Wa3W28sd+0RAjvfuUafSh6O4lOT3LSR3jfTxCkJx+eDVPGGpurtXIZYQpcy8zcJSkZILis/RIP0FaMDJiyT4pG7dVn57qjJhfD6L2dX2XZauhybDsZpGDhTf29NkXSSMY5BIShoH8knl71JnZi250redDuahvdqi3SXJkuspTJHNDSEYGAnwyTkk+PhXZb77aOam24gQLChBnWMJ+BaJx3rYQEKbz4ZAII/ED1qu2hNydZbZKmWmMy22hbhU7CuLCh3TmMFQGQQegz5HAqlr352G5sLqfqJO9K4sGNMDILbS+b9aat+jNzZtrsfJiKG2pUdHIksFQzxB8cAjI8wCK77tOxXrrozQutHEkvyYSWJSsea20uD9ef1rgtN6f1fvBrhydIDz/wAU6kz7iW+LLDY6YHlkJ6JSOv6mrK7+6aYmbIXKDEa+W1x25EZPiUhnH/pyqcjJEE2O1zrcNj+y5ii8VkrgKB4UF6/1b8X2c9D2NLuXXHnEvjPXhGJSn9VJ+lclIt9w261hpe6v8gssQ7qMjGErOVJ9gCK02k4T+pdR2PTgdLjciYlltGeiErWCsj2BPtU99suwt/Y+n7+w0EtxlqgOY8AhSeSPbKVD3rQXx40rcb2eST9eFXTpmOm/1pcVvsv9su0A1Zoii62sw4DeP4VAKUf75NXBhtNsRWmGU8W2kBCE+gAwKp32Ybe9qLeRm5y1mR9nRlylrV/EEhtGf7X6VcdsYQBXi9WIYWQA7NH3/AvRwAXB0h9yv1SlK8legleE2OiVHcYcGW3EFCxnGUkYI+le9KIo3GyW1oAA0lHwOn84e/8AqvSJsxtnFkokM6ShKWg5AdW44nP+6pRB96kTA9BTA9Kv83kVWs/9KqEEY/xCxXLdAdQ2l2DFWG08UBTSSEj0HToK9mI8dhsNMMNNIByEoQEj6CvSlU2VZQWK7bbc84px23xHFqOVKUykk/mcVUvthJSjdCKlKQlKbS0AAMADm5Vvq4vWG2mjNXXZN01DZxNlpaDKVl5xGEAkgYSoDxJrf03NZiTiR9kLNlwGaPS1fjaeFDm7SaYamRI8lv7NZPB5sLH3B5EV2UaNHjMJZjsNMtJ8ENoCUj2FY1jtcSzWqNa7e0GYcVoNMN8irigdAMnqazqxSP1vcfYlaGN0gBfOKfQfSsNyz2lyT8Uu1wVPg571UdJX9cZrNpXIJHCkgFfMD0FYz1ttzzinHoEVxavvKUykk/mcVlUoCRwhAPKAAAADAHlXxSEqBCkggjByPGvtKhSsdiFDYcLjEOO0sjBUhsJJ9wK0UPQ2k4OondQw9PW5i6OEqVJQyAvJ8SPIE+ZHU10tMCug9w4K5LWnkL8FsKa4EDGMEY8a5zTuhtJafur1ys2nbbBlu8ubrTICuviAf6I/AYrpqVAJAoHZSWg8hfOIKcEAj0xWLOtltnkGdb4kojwLzKV4+orLpUDbhSRa82WGWGg0yy220nwQhIAHsK/S20LQW1oSpChgpIyCK/VKIsRq225lxLjVviNrT1SpLKQR74r2kRo8hvu5DDTqM54rQFDP5GvWlTZu1FBeEaFDjKKo0RhlShgltsJJHtXvSlQTamqSlKURKUpREpSlESlKURKYHpSlESlKURKUpREpSlESlKURKUpREpSlESlKURKUpRF//Z"

# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="ExpertNow – CV Builder", page_icon="📄", layout="wide")

# Brand colors extracted from ExpertNow logo
# Navy: #1e3a5f  |  Red accent: #e8251f  |  Light navy: #2a5298
# Background: clean white/off-white  |  Subtle grays for structure

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&family=DM+Serif+Display:ital@0;1&display=swap');

/* ── Reset & Base ── */
html, body, [data-testid="stAppViewContainer"] {
    background-color: #f0f2f5 !important;
    font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stAppViewContainer"] > .main {
    background-color: #f0f2f5 !important;
}

/* ── Top Banner ── */
.top-banner {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0 32px;
    background: #ffffff;
    border-bottom: 4px solid #e8251f;
    margin-bottom: 28px;
    height: 72px;
    box-shadow: 0 2px 12px rgba(30,58,95,0.10);
}
.top-banner-left {
    display: flex;
    align-items: center;
    gap: 20px;
}
.top-banner img { height: 40px; }
.banner-divider {
    width: 1px; height: 32px;
    background: #d0d8e4;
}
.banner-title {
    font-family: 'DM Sans', sans-serif;
    font-size: 18px;
    font-weight: 700;
    color: #1e3a5f;
    letter-spacing: 0.3px;
}
.banner-subtitle {
    font-size: 12px;
    color: #7a8fa6;
    margin-top: 1px;
    font-weight: 400;
}
.banner-badge {
    background: linear-gradient(135deg, #1e3a5f 0%, #2a5298 100%);
    color: #fff;
    font-size: 11px;
    font-weight: 600;
    padding: 5px 14px;
    border-radius: 20px;
    letter-spacing: 0.8px;
    text-transform: uppercase;
}

/* ── Section Headers ── */
.section-header {
    display: flex;
    align-items: center;
    gap: 10px;
    margin: 28px 0 10px;
}
.section-header-line {
    flex: 1;
    height: 1px;
    background: linear-gradient(90deg, #d0d8e4 0%, transparent 100%);
}
.section-header-text {
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: #1e3a5f;
    background: #fff;
    padding: 5px 14px;
    border-radius: 3px;
    border-left: 3px solid #e8251f;
    box-shadow: 0 1px 4px rgba(30,58,95,0.08);
    white-space: nowrap;
}

/* ── Cards / Expanders ── */
[data-testid="stExpander"] {
    border: 1px solid #dde6ef !important;
    border-radius: 8px !important;
    background: #ffffff !important;
    margin-bottom: 8px !important;
    box-shadow: 0 1px 4px rgba(30,58,95,0.06) !important;
    transition: box-shadow 0.2s ease;
}
[data-testid="stExpander"]:hover {
    box-shadow: 0 3px 12px rgba(30,58,95,0.12) !important;
}
[data-testid="stExpanderDetails"] {
    background: #fff !important;
    padding: 12px 16px !important;
    border-top: 1px solid #eef1f6 !important;
}

/* ── Inputs ── */
[data-testid="stTextInput"] > div > div > input,
[data-testid="stTextArea"] > div > div > textarea,
[data-testid="stNumberInput"] > div > div > input,
[data-testid="stSelectbox"] > div > div {
    border: 1px solid #d0d8e4 !important;
    border-radius: 6px !important;
    background: #f8fafc !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 13.5px !important;
    color: #1e2d3d !important;
    transition: border-color 0.15s, box-shadow 0.15s;
}
[data-testid="stTextInput"] > div > div > input:focus,
[data-testid="stTextArea"] > div > div > textarea:focus {
    border-color: #1e3a5f !important;
    box-shadow: 0 0 0 3px rgba(30,58,95,0.10) !important;
    background: #ffffff !important;
}

/* Input Labels */
[data-testid="stTextInput"] label,
[data-testid="stTextArea"] label,
[data-testid="stNumberInput"] label,
[data-testid="stSelectbox"] label {
    font-size: 11.5px !important;
    font-weight: 600 !important;
    color: #4a6080 !important;
    letter-spacing: 0.3px !important;
    text-transform: uppercase !important;
}

/* ── Buttons ── */
.stButton > button {
    background: #1e3a5f !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 6px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    padding: 8px 18px !important;
    letter-spacing: 0.3px !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 2px 6px rgba(30,58,95,0.20) !important;
}
.stButton > button:hover {
    background: #e8251f !important;
    box-shadow: 0 4px 14px rgba(232,37,31,0.30) !important;
    transform: translateY(-1px) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}

/* ── Generate Button (primary) ── */
.generate-btn .stButton > button {
    background: linear-gradient(135deg, #e8251f 0%, #c01c17 100%) !important;
    font-size: 14px !important;
    padding: 12px 24px !important;
    letter-spacing: 0.5px !important;
    box-shadow: 0 4px 16px rgba(232,37,31,0.35) !important;
}
.generate-btn .stButton > button:hover {
    background: linear-gradient(135deg, #c01c17 0%, #a01512 100%) !important;
    box-shadow: 0 6px 20px rgba(232,37,31,0.45) !important;
}

/* ── Download button ── */
[data-testid="stDownloadButton"] > button {
    background: linear-gradient(135deg, #1e3a5f 0%, #2a5298 100%) !important;
    color: #fff !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    border: none !important;
    transition: all 0.2s !important;
}
[data-testid="stDownloadButton"] > button:hover {
    box-shadow: 0 4px 14px rgba(30,58,95,0.35) !important;
    transform: translateY(-1px) !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid #dde6ef !important;
}
[data-testid="stSidebar"] .stMarkdown h3 {
    color: #1e3a5f !important;
    font-size: 13px !important;
    letter-spacing: 0.5px !important;
    text-transform: uppercase !important;
}

/* ── Info / Success / Error ── */
[data-testid="stAlert"] {
    border-radius: 6px !important;
}

/* ── Divider ── */
hr {
    border: none !important;
    border-top: 1px solid #dde6ef !important;
    margin: 24px 0 !important;
}

/* ── JSON Preview ── */
[data-testid="stJson"] {
    border-radius: 8px !important;
    border: 1px solid #dde6ef !important;
}

/* ── Column gaps ── */
[data-testid="column"] {
    padding: 0 6px !important;
}

/* ── Number input ── */
[data-testid="stNumberInput"] button {
    background: #f0f2f5 !important;
    border-color: #d0d8e4 !important;
    color: #1e3a5f !important;
}

/* Subtle page grid lines */
[data-testid="stAppViewContainer"] > .main > .block-container {
    padding-top: 0 !important;
    max-width: 1200px !important;
}
</style>
""", unsafe_allow_html=True)

# ── Logo bytes for DOCX
try:
    LOGO_BYTES = base64.b64decode(_LOGO_B64)
except Exception:
    LOGO_BYTES = None

_logo_uri = "data:image/jpeg;base64," + _LOGO_B64 if _LOGO_B64 else ""
_img_tag = f"<img src='{_logo_uri}' alt='ExpertNow'>" if _logo_uri else ""

st.markdown(f"""
<div class="top-banner">
  <div class="top-banner-left">
    {_img_tag}
    <div class="banner-divider"></div>
    <div>
      <div class="banner-title">CV Builder</div>
      <div class="banner-subtitle">Créez votre CV professionnel au format ExpertNow</div>
    </div>
  </div>
  <div class="banner-badge">✦ Pro Tool</div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────────────────────────────────────
def _default():
    return {
        "full_name": "", "title": "", "years_experience": 0,
        "certifications_summary": "",
        "certifications": [{"year": "", "items": ""}],
        "education":      [{"years": "", "degree": "", "institution": "", "mention": ""}],
        "languages":      [{"name": "", "level": ""}],
        "skills":         [{"category": "", "items": ""}],
        "experiences":    [{"company": "", "role": "", "period": "",
                            "clients": "", "missions": [""], "technologies": ""}],
    }

if "cv" not in st.session_state:
    st.session_state.cv = _default()

def _add(key, tpl):
    st.session_state.cv[key].append(copy.deepcopy(tpl))

def _rm(key, i):
    if len(st.session_state.cv[key]) > 1:
        st.session_state.cv[key].pop(i)

def sh(label, icon=""):
    st.markdown(f"""
    <div class="section-header">
        <div class="section-header-text">{icon + " " if icon else ""}{label}</div>
        <div class="section-header-line"></div>
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Dropdown option lists
# ─────────────────────────────────────────────────────────────────────────────
CURRENT_YEAR    = 2025
YEAR_OPTIONS    = [""] + [str(y) for y in range(CURRENT_YEAR, 1979, -1)]
LANGUAGE_LEVELS = ["", "Notions", "Intermédiaire", "Courant", "Avancé", "Bilingue", "Langue maternelle"]
MENTION_OPTIONS = ["", "Très Bien", "Bien", "Assez Bien", "Passable", "Avec Mention"]

# ─────────────────────────────────────────────────────────────────────────────
# Sidebar
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    if _logo_uri:
        st.markdown(f'<img src="{_logo_uri}" style="width:150px;margin:16px 0 8px;">', unsafe_allow_html=True)

    st.markdown("""
    <div style="background:linear-gradient(135deg,#1e3a5f,#2a5298);border-radius:8px;padding:12px 14px;margin:8px 0 16px;color:white;">
        <div style="font-size:11px;font-weight:700;letter-spacing:1px;text-transform:uppercase;opacity:0.7;margin-bottom:4px;">Navigation</div>
        <div style="font-size:13px;font-weight:500;">Remplissez chaque section du formulaire puis générez votre CV.</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🔄 Réinitialiser le formulaire", use_container_width=True):
        st.session_state.cv = _default()
        st.rerun()

    st.markdown("### 💾 Import / Export")
    st.download_button(
        "⬇️ Exporter en JSON",
        data=json.dumps(st.session_state.cv, ensure_ascii=False, indent=2),
        file_name="cv_data.json", mime="application/json",
        use_container_width=True,
    )

    up = st.file_uploader("⬆️ Importer JSON", type="json")
    if up is not None:
        file_id = f"{up.name}_{up.size}"
        if st.session_state.get("_last_upload_id") != file_id:
            try:
                loaded = json.load(up)
                loaded["years_experience"] = int(loaded.get("years_experience") or 0)
                st.session_state.cv = loaded
                st.session_state["_last_upload_id"] = file_id
                st.success("✅ Données importées !")
            except Exception as e:
                st.error(f"JSON invalide : {e}")
    else:
        st.session_state.pop("_last_upload_id", None)

    st.markdown("---")
    st.markdown("""
    <div style="font-size:11px;color:#7a8fa6;line-height:1.6;padding:4px 0;">
        <strong style="color:#1e3a5f;">ExpertNow CV Builder</strong><br>
        Outil de création de CV au format standardisé ExpertNow.
    </div>
    """, unsafe_allow_html=True)

cv = st.session_state.cv

# ─────────────────────────────────────────────────────────────────────────────
# Form
# ─────────────────────────────────────────────────────────────────────────────

# 1 · Identity
sh("Identité du Consultant", "👤")
c1, c2, c3 = st.columns([2, 2, 1])
with c1:
    cv["full_name"] = st.text_input("Nom complet", cv["full_name"])
with c2:
    cv["title"] = st.text_input("Titre du poste", cv["title"])
with c3:
    cv["years_experience"] = st.number_input(
        "Années d'exp.", min_value=0, max_value=60,
        value=int(cv.get("years_experience") or 0), step=1,
    )
cv["certifications_summary"] = st.text_input(
    "Résumé certifications (en-tête)", cv["certifications_summary"]
)

# 2 · Certifications
sh("Certifications", "🏅")
for i, cert in enumerate(cv["certifications"]):
    with st.expander(f"Groupe de certifications {i+1}", expanded=(i == 0)):
        c1, c2, c3 = st.columns([1, 5, 1])
        with c1:
            cur = cert.get("year", "")
            idx = YEAR_OPTIONS.index(cur) if cur in YEAR_OPTIONS else 0
            cert["year"] = st.selectbox("Année", YEAR_OPTIONS, index=idx, key=f"cy{i}")
        with c2:
            cert["items"] = st.text_area("Éléments (un par ligne)", cert["items"], height=80, key=f"ci{i}")
        with c3:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🗑️", key=f"rc{i}"):
                _rm("certifications", i); st.rerun()
if st.button("➕ Ajouter un groupe"):
    _add("certifications", {"year": "", "items": ""}); st.rerun()

# 3 · Education
sh("Formation", "🎓")
for i, edu in enumerate(cv["education"]):
    with st.expander(f"Formation {i+1} — {edu.get('degree') or 'Nouveau diplôme'}", expanded=(i == 0)):
        c1, c2 = st.columns([1, 3])
        with c1:
            edu["years"] = st.text_input("Années (ex: 2018-2021)", edu["years"], key=f"ey{i}")
        with c2:
            edu["degree"] = st.text_input("Diplôme", edu["degree"], key=f"ed{i}")
        c3, c4 = st.columns([3, 1])
        with c3:
            edu["institution"] = st.text_input("Établissement", edu["institution"], key=f"ei{i}")
        with c4:
            cur = edu.get("mention", "")
            idx = MENTION_OPTIONS.index(cur) if cur in MENTION_OPTIONS else 0
            edu["mention"] = st.selectbox("Mention", MENTION_OPTIONS, index=idx, key=f"em{i}")
        if st.button("🗑️ Supprimer cette formation", key=f"re{i}"):
            _rm("education", i); st.rerun()
if st.button("➕ Ajouter une formation"):
    _add("education", {"years": "", "degree": "", "institution": "", "mention": ""}); st.rerun()

# 4 · Languages
sh("Langues", "🌍")
cols = st.columns(4)
for i, lang in enumerate(cv["languages"]):
    with cols[i % 4]:
        lang["name"] = st.text_input("Langue", lang["name"], key=f"ln{i}")
        cur = lang.get("level", "")
        idx = LANGUAGE_LEVELS.index(cur) if cur in LANGUAGE_LEVELS else 0
        lang["level"] = st.selectbox("Niveau", LANGUAGE_LEVELS, index=idx, key=f"ll{i}")
        if st.button("🗑️", key=f"rl{i}"):
            _rm("languages", i); st.rerun()
if st.button("➕ Ajouter une langue"):
    _add("languages", {"name": "", "level": ""}); st.rerun()

# 5 · Skills
sh("Compétences Techniques", "🛠️")
for i, sk in enumerate(cv["skills"]):
    c1, c2, c3 = st.columns([2, 5, 1])
    with c1:
        sk["category"] = st.text_input("Catégorie", sk["category"], key=f"sc{i}")
    with c2:
        sk["items"] = st.text_input("Outils / technologies", sk["items"], key=f"si{i}")
    with c3:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🗑️", key=f"rs{i}"):
            _rm("skills", i); st.rerun()
if st.button("➕ Ajouter une catégorie"):
    _add("skills", {"category": "", "items": ""}); st.rerun()

# 6 · Experiences
sh("Expériences Professionnelles", "💼")
for i, exp in enumerate(cv["experiences"]):
    title_label = f"📌 {exp['company']}" if exp['company'] else f"Expérience {i+1}"
    if exp.get('role'): title_label += f"  ·  {exp['role']}"
    with st.expander(title_label, expanded=(i == 0)):
        c1, c2 = st.columns(2)
        with c1:
            exp["company"] = st.text_input("Entreprise", exp["company"], key=f"eco{i}")
            exp["role"]    = st.text_input("Rôle / titre", exp["role"], key=f"ero{i}")
        with c2:
            exp["period"]  = st.text_input("Période (ex: 2022 - Présent)", exp["period"], key=f"epe{i}")
            exp["clients"] = st.text_input("Clients (optionnel)", exp["clients"], key=f"ecl{i}")
        st.markdown("**Missions**")
        for j, m in enumerate(exp["missions"]):
            mc1, mc2 = st.columns([10, 1])
            with mc1:
                exp["missions"][j] = st.text_area(
                    f"Mission {j+1}", m, height=60,
                    key=f"em{i}_{j}", label_visibility="collapsed"
                )
            with mc2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🗑️", key=f"rm{i}_{j}"):
                    if len(exp["missions"]) > 1:
                        exp["missions"].pop(j); st.rerun()
        if st.button("➕ Ajouter une mission", key=f"am{i}"):
            exp["missions"].append(""); st.rerun()
        exp["technologies"] = st.text_area(
            "Technologies & mots-clés", exp["technologies"], height=60, key=f"et{i}"
        )
        if st.button("🗑️ Supprimer cette expérience", key=f"rx{i}"):
            _rm("experiences", i); st.rerun()
if st.button("➕ Ajouter une expérience"):
    _add("experiences", {"company": "", "role": "", "period": "", "clients": "",
                         "missions": [""], "technologies": ""})
    st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# DOCX generation
# ─────────────────────────────────────────────────────────────────────────────
def _rule(p, color="1e3a5f", sz=8):
    pPr = p._p.get_or_add_pPr()
    bd  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    bd.append(bot); pPr.append(bd)

def _sh(doc, text, col="1e3a5f"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(int(col[0:2], 16), int(col[2:4], 16), int(col[4:6], 16))
    _rule(p, color=col)

def build_docx(cv):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.7)
        sec.left_margin = sec.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    if LOGO_BYTES:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_after = Pt(2)
        lp.add_run().add_picture(io.BytesIO(LOGO_BYTES), width=Inches(1.7))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cv["full_name"].upper())
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    if cv["title"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cv["title"]); r.bold = True; r.font.size = Pt(13)

    yrs = cv.get("years_experience", 0)
    if yrs:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{yrs} ans d'expérience"); r.bold = True; r.font.size = Pt(11)

    if cv["certifications_summary"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"({cv['certifications_summary']})")
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    rp = doc.add_paragraph()
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after  = Pt(4)
    _rule(rp, color="e8251f", sz=14)

    if any(c["year"] or c["items"] for c in cv["certifications"]):
        _sh(doc, "Certifications")
        for cert in cv["certifications"]:
            if not cert["year"] and not cert["items"]: continue
            for idx, item in enumerate([l.strip() for l in cert["items"].splitlines() if l.strip()]):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(1)
                if idx == 0 and cert["year"]:
                    r = p.add_run(f"{cert['year']} : "); r.bold = True
                p.add_run(item)

    if any(e["degree"] for e in cv["education"]):
        _sh(doc, "Education")
        for edu in cv["education"]:
            if not edu["degree"]: continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            if edu["years"]:
                r = p.add_run(f"{edu['years']} : "); r.bold = True
            r = p.add_run(edu["degree"]); r.bold = True
            if edu["institution"]: p.add_run(f"\n    {edu['institution']}")
            if edu["mention"]:     p.add_run(f", {edu['mention']}")

    if any(l["name"] for l in cv["languages"]):
        _sh(doc, "Langues")
        for lang in cv["languages"]:
            if not lang["name"]: continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(lang["name"]); r.bold = True
            if lang["level"]: p.add_run(f" : {lang['level']}")

    if any(s["category"] for s in cv["skills"]):
        _sh(doc, "Compétences Techniques")
        for sk in cv["skills"]:
            if not sk["category"]: continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(f"{sk['category']} : "); r.bold = True
            p.add_run(sk["items"])

    if any(e["company"] for e in cv["experiences"]):
        _sh(doc, "Expériences Professionnelles")
        for exp in cv["experiences"]:
            if not exp["company"]: continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(exp["company"]); r.bold = True
            r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
            if exp["role"]:
                sep = p.add_run(" | "); sep.bold = True
                r = p.add_run(exp["role"]); r.bold = True
            if exp["period"]:
                p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1)
                r = p2.add_run(exp["period"]); r.italic = True; r.font.size = Pt(9)
            if exp["clients"]:
                p3 = doc.add_paragraph(); p3.paragraph_format.space_after = Pt(2)
                p3.add_run(f"Clients : {exp['clients']}").font.size = Pt(9)
            ml = doc.add_paragraph(); ml.paragraph_format.space_after = Pt(1)
            ml.add_run("Missions :").bold = True
            for mission in exp["missions"]:
                if not mission.strip(): continue
                mp = doc.add_paragraph(style="List Bullet")
                mp.paragraph_format.left_indent  = Inches(0.25)
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after  = Pt(1)
                mp.add_run(mission.strip())
            if exp["technologies"]:
                tp = doc.add_paragraph()
                tp.paragraph_format.space_before = Pt(3)
                tp.paragraph_format.space_after  = Pt(2)
                r = tp.add_run("Technologies et mots-clés : "); r.bold = True
                tp.add_run(exp["technologies"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# Generate button
# ─────────────────────────────────────────────────────────────────────────────
st.divider()

st.markdown("""
<div style="background:linear-gradient(135deg,#f8fafc,#eef2f8);border:1px solid #dde6ef;border-radius:10px;padding:20px 24px;margin-bottom:16px;">
    <div style="font-size:13px;font-weight:700;color:#1e3a5f;letter-spacing:0.5px;text-transform:uppercase;margin-bottom:4px;">
        ✦ Génération du CV
    </div>
    <div style="font-size:13px;color:#7a8fa6;">
        Vérifiez vos informations puis cliquez sur le bouton pour générer votre CV au format Word ExpertNow.
    </div>
</div>
""", unsafe_allow_html=True)

col_gen, col_prev = st.columns([2, 3])

with col_gen:
    st.markdown('<div class="generate-btn">', unsafe_allow_html=True)
    if st.button("🚀 Générer le CV (.docx)", use_container_width=True):
        if not cv["full_name"]:
            st.error("⚠️ Veuillez saisir un nom complet avant de générer.")
        else:
            with st.spinner("Génération en cours…"):
                try:
                    docx_bytes = build_docx(cv)
                    st.success("✅ CV généré avec succès !")
                    st.download_button(
                        label="⬇️ Télécharger CV.docx",
                        data=docx_bytes,
                        file_name=f"{cv['full_name'].replace(' ', '_')}_CV.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Erreur lors de la génération : {e}")
    st.markdown('</div>', unsafe_allow_html=True)

with col_prev:
    with st.expander("👁️ Prévisualisation des données (JSON)"):
        st.json(cv)